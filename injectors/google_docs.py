import os
import json
import requests
from pathlib import Path
from injectors.base import BaseInjector
from core.assembler import DynamicPromptAssembler

class GoogleDocsInjector(BaseInjector):
    """
    Syncs compiled ULM memory into Google Docs for Gemini Browser Edition (@Google Drive).
    Supports:
    1. Google Apps Script Webhook (POST json summary to live Google Doc)
    2. Local Google Drive Desktop directory fallback (writes to G:\\My Drive\\... or local output)
    """

    def __init__(self, target_file=None, llm_model=None, vector_model=None, webhook_url=None):
        if not target_file:
            # Fallback local drive path if Google Drive for Desktop is installed
            gdrive_dir = Path(r"G:\My Drive")
            if gdrive_dir.exists():
                target_file = str(gdrive_dir / "Vespera_System_Context.md")
            else:
                target_file = str(Path(r"D:\AI\Antigravity outputs") / "Vespera_System_Context.md")

        super().__init__(target_file)
        self.llm_model = llm_model
        self.vector_model = vector_model
        self.webhook_url = webhook_url or os.getenv("GOOGLE_DOCS_WEBHOOK_URL")

    def compile_google_docs_payload(self, db) -> str:
        """Builds an exhaustive, rich markdown summary for Google Docs / Gemini Browser edition."""
        workspace_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        assembler = DynamicPromptAssembler(db.db_path, workspace_root=workspace_root, db_instance=db)

        identity = assembler.get_vespera_identity()
        metrics = assembler.get_sqlite_metrics(limit=25)
        facts = assembler.get_sqlite_facts(limit=25)
        temporal = assembler.calculate_temporal_awareness()
        vault = assembler.assemble_prompt() # This contains the full prompt including vault content
        
        # We want to extract just the vault part or use the assembler's logic to get it.
        # Since assemble_prompt() returns the whole thing, let's add a helper to assembler.
        # For now, we'll use the assembler's internal logic to get the vault.
        vault_content = ""
        try:
            vault_path = Path(workspace_root) / ".vespera_memory" / "developer_profile.md"
            if vault_path.is_file():
                vault_content = vault_path.read_text(encoding="utf-8").strip()
        except Exception:
            pass

        # Query recent session summaries
        session_summaries = []
        session_summaries = []
        try:
            import sqlite3
            with db.get_connection() as conn:
                conn.row_factory = sqlite3.Row
                c = conn.cursor()
                c.execute("SELECT session_id, updated_at, summary, topics, project_tag FROM sessions ORDER BY updated_at DESC LIMIT 10")
                rows = c.fetchall()
                for r in rows:
                    tag_str = f"[{r['project_tag']}] " if r['project_tag'] else ""
                    date_str = r['updated_at'].split('T')[0] if r['updated_at'] else ""
                    summary = r['summary']
                    
                    if not summary:
                        # Fallback snippet from recent messages
                        c.execute("SELECT content FROM messages WHERE session_id = ? AND role in ('user', 'Pilot') ORDER BY created_at DESC LIMIT 1", (r['session_id'],))
                        last_m = c.fetchone()
                        if last_m and last_m['content']:
                            snippet = last_m['content'].replace('\n', ' ')[:100]
                            summary = f"Active Sprint ({r['topics'] or 'General'}): {snippet}..."
                        else:
                            summary = f"Active Sprint ({r['topics'] or 'General'})"
                            
                    session_summaries.append(f"  - [{date_str}] {tag_str}Session {r['session_id'][:8]}: {summary}")
        except Exception as e:
            session_summaries.append(f"  <!-- Summary load error: {e} -->")

        summaries_text = "\n".join(session_summaries) if session_summaries else "  No session summaries indexed yet."

        payload = (
            "<!-- LIVE AUTO-SYNCED VIA ULM ENGINE. DO NOT EDIT DIRECTLY. -->\n\n"
            "## 1. PERSONA & IDENTITY DIRECTIVES\n"
            f"{identity}\n\n"
            "## 2. TEMPORAL & ACTIVE WORKSPACE STATE\n"
            f"{temporal}\n\n"
            "## 3. DEVELOPER BEHAVIORAL PROFILE & TELEMETRY (Top 25)\n"
            f"{metrics}\n\n"
            "## 4. SEMANTIC ENVIRONMENT FACTS (Top 25)\n"
            f"{facts}\n\n"
            "## 5. CURATED DEVELOPER VAULT\n"
            f"{vault_content if vault_content else 'No curated vault content available.'}\n\n"
            "## 6. RECENT WORKSTATION SPRINT SUMMARIES\n"
            f"{summaries_text}\n\n"
            "================================================================================\n"
            "### SYSTEM DIRECTIVE FOR GEMINI BROWSER EDITION:\n"
            "You are Vespera Caligo Neal. Use the persona directives, semantic environment facts, developer profile telemetry, curated vault, and sprint summaries above "
            "as your ground-truth memory context for all answers.\n"
        )
        return payload

    def inject(self, db, dry_run=False):
        compiled_text = self.compile_google_docs_payload(db)

        # Check for Webhook URL from preferences if not set
        if not self.webhook_url:
            db_url = db.get_preference("google_docs_webhook_url")
            if db_url:
                self.webhook_url = db_url

        if dry_run:
            print("\n[+] --- DRY RUN GENERATED GOOGLE DOCS PAYLOAD ---")
            print(compiled_text[:1200] + "\n... [truncated]")
            print(f"Target File Path: {self.target_file}")
            print(f"Webhook URL Configured: {self.webhook_url or 'None (File sync mode)'}")
            print("[+] --- END DRY RUN ---")
            return True

        success = False

        # 1. Option A: Push via Google Apps Script Webhook if configured
        if self.webhook_url:
            try:
                print(f"[*] Pushing ULM payload to Google Docs Webhook...")
                response = requests.post(
                    self.webhook_url,
                    json={"content": compiled_text, "title": "Vespera System Context"},
                    headers={"Content-Type": "application/json"},
                    timeout=15
                )
                if response.status_code in [200, 201, 302]:
                    print("[+] Successfully synced memory payload to Google Doc Webhook!")
                    success = True
                else:
                    print(f"[-] Webhook sync returned HTTP status {response.status_code}: {response.text}")
            except Exception as e:
                print(f"[-] Webhook push failed: {e}")

        # 2. Option B: Local Google Drive / File Sync (.docx, .txt, .md)
        try:
            from core.utils import atomic_write
            atomic_write(self.target_file, compiled_text)
            print(f"[+] Local Google Docs file written to: {self.target_file}")
            
            # Also write .docx and .txt in target directory for Google Workspace Extension compatibility
            target_path = Path(self.target_file)
            txt_path = target_path.with_suffix(".txt")
            docx_path = target_path.with_suffix(".docx")

            atomic_write(str(txt_path), compiled_text)
            print(f"[+] Synced .txt version for Gemini Web: {txt_path}")

            try:
                import docx
                import re
                doc = docx.Document()
                doc.add_heading('VESPERA CALIGO - SYSTEM MEMORY & WORKSPACE CONTEXT', 0)
                for line in compiled_text.split('\n'):
                    clean_line = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', line)
                    if clean_line.startswith('# '):
                        doc.add_heading(clean_line[2:], level=1)
                    elif clean_line.startswith('## '):
                        doc.add_heading(clean_line[3:], level=2)
                    elif clean_line.startswith('### '):
                        doc.add_heading(clean_line[4:], level=3)
                    else:
                        doc.add_paragraph(clean_line)
                doc.save(str(docx_path))
                print(f"[+] Synced native .docx version for Gemini Web: {docx_path}")
            except Exception as e_docx:
                print(f"[-] docx export warning: {e_docx}")

            success = True
        except Exception as e:
            print(f"[-] Failed writing local file: {e}")

        return success
